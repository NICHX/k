# -*- coding: utf-8 -*-
import xlwt
from DrissionPage import Chromium
from DrissionPage.errors import ElementNotFoundError
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt
import os
import sys



def download_ques_enterprise(ID, delay, begin, file_format, anl_switch):
    try:
        os.mkdir(rf'.\{ID}')
        print('目录创建成功')
    except FileExistsError:
        print('目录已存在', flush=True)
        pass

    doc = Document()
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = Pt(11)

    wb = xlwt.Workbook(encoding='utf-8')  # 新建一个excel文件
    ws1 = wb.add_sheet('sheet1')  # 添加一个新表，名字为begin
    ws1.write(0, 0, '序号')
    ws1.write(0, 1, '题目')
    ws1.write(0, 2, 'A')
    ws1.write(0, 3, 'B')
    ws1.write(0, 4, 'C')
    ws1.write(0, 5, 'D')
    ws1.write(0, 6, 'E')
    ws1.write(0, 7, 'F')
    ws1.write(0, 8, 'G')
    ws1.write(0, 9, 'H')
    ws1.write(0, 10, '正确答案')
    ws1.write(0, 11, '解析')
    login_url = f'https://s.kaoshibao.com/student'
    url = f'https://s.kaoshibao.com/online/?paperId={ID}'
    browser = Chromium()
    tab = browser.latest_tab
    if tab.url != 'https://s.kaoshibao.com/sctk/':
        tab.get(login_url)
        # 获取文本框元素对象
        tab.wait.url_change('https://s.kaoshibao.com/sctk/', timeout=30)
        tab.wait(1)
        tab = browser.new_tab(url)
    elif tab.url == 'https://s.kaoshibao.com/sctk/':
        tab = browser.new_tab(url)

    tab.wait.eles_loaded('xpath://*[@id="body"]/div/div[1]/div[2]/div[1]/div[1]/div/div[1]/div/span[2]')
    number = tab.ele('xpath://*[@id="body"]/div/div[1]/div[2]/div[1]/div[1]/div/div[1]/div/span[2]').text[2:-1]
    try:
        tab.ele(f'tag:span@text():{begin}').click()
    except Exception as e:
        print(e)
    # 打开背题模式
    try:
        tab.wait(3)
        button_off = tab.ele('xpath://*[@id="body"]/div/div[1]/div[2]/div[2]/div[3]/p[2]/span[2]/div',timeout=0.5)
        auto_bext_button = tab.ele('xpath://*[@id="body"]/div/div[1]/div[2]/div[2]/div[3]/p[1]/span[2]/div',
                                     timeout=0.1)
        try:
            a = auto_bext_button.attr('aria-checked')
            if a == 'true':
                tab.ele('xpath://*[@id="body"]/div/div[1]/div[2]/div[2]/div[3]/p[1]/span[2]/div/input').click()
                print('关闭答对自动下一题', flush=True)
            else:
                print('答对自动下一题已关闭', flush=True)
        except Exception as e:
            print(e)
            sys.exit(0)
        try:
            a = button_off.attr('aria-checked')
            if a is None:
                tab.ele('xpath://*[@id="body"]/div/div[1]/div[2]/div[2]/div[3]/p[2]/span[2]/div/input').click()
                print('点击背题模式按钮', flush=True)
            else:
                print('背题模式已打开或已禁用', flush=True)
        except Exception as e:
            raise Exception(ElementNotFoundError)
    except ElementNotFoundError:
        print('背题模式已打开或已禁用', flush=True)
        tab.wait(0.3)

    for i in range(int(begin) - 1, int(number)):
        try:
            title = f"{tab.ele('@class=qusetion-box').text}".replace('\n', '')
            doc.add_paragraph(f'{i + 1}.{title}')
            try:
                ques_img = tab.ele('@class=qusetion-box').ele('tag:img',timeout=0.1).get_screenshot(path=rf'.\{ID}\imgs\title',name=f'ques{i + 1}.png')
                doc.add_picture(ques_img)
            except Exception as e:
                pass
            topic = tab.ele('@class=topic-type').text
            option = ''
            if topic == '单选题':
                try:
                    tab.ele('@@class^before-icon@@text()=A').click()
                except Exception as e:
                    print(e)
                options = tab.eles('@class^option')
                for j in options:
                    try:
                        x = j.s_ele('@class^before-icon')
                        option_img = j.ele('tag:img',timeout=0.1).get_screenshot(path=rf'.\{ID}\imgs\option',
                                                                                  name=f'ques{i + 1}-option-{x.text}.png')
                        para = doc.add_paragraph()
                        list_j = list(j.text)
                        list_j.insert(1, '.')
                        str_j = ''.join(list_j)
                        run = para.add_run(str_j)  # 添加选项文本
                        run.add_picture(option_img)
                    except Exception as e:
                        list_j = list(j.text)
                        list_j.insert(1, '.')
                        str_j = ''.join(list_j)
                        doc.add_paragraph(str_j)
                        option += str_j + "\n"
                answer = tab.s_ele('@class=right-ans').text.replace('\u2003', ':')
            elif topic == '判断题':
                try:
                    tab.ele('@@class^before-icon@@text()=A').click()
                except Exception as e:
                    print(e)
                options = tab.ele('@class^select-left').children('@class^option')
                for j in options:
                    list_j = list(j.text)
                    list_j.insert(1, '.')
                    str_j = ''.join(list_j)
                    doc.add_paragraph(str_j)
                    option += str_j + "\n"
                answer = tab.s_ele('@class=right-ans').text.replace('\u2003', ':')
            elif topic == '多选题':
                if tab.s_ele('xpath://*[@id="body"]/div/div[1]/div[2]/div[2]/div[3]/p[2]/span[2]/div').attr(
                        'aria-checked') is None:
                    try:
                        tab.ele('@@class^before-icon@@text()=A').click()
                        tab.wait(0.1)
                        tab.ele('xpath://*[@id="body"]/div/div[1]/div[2]/div[1]/div[1]/div/div[3]/button').click()
                    except Exception as e:
                        print(e)
                else:
                    pass
                options = tab.eles('@class^option')
                for j in options:
                    try:
                        x = j.s_ele('@class^before-icon')
                        option_img = j.ele('tag:img',timeout=0.1).get_screenshot(path=rf'.\{ID}\imgs\option',
                                                                                  name=f'ques{i + 1}-option-{x.text}.png')
                        para = doc.add_paragraph()
                        list_j = list(j.text)
                        list_j.insert(1, '.')
                        str_j = ''.join(list_j)
                        run = para.add_run(str_j)  # 添加选项文本
                        run.add_picture(option_img)
                    except Exception as e:
                        list_j = list(j.text)
                        list_j.insert(1, '.')
                        str_j = ''.join(list_j)
                        doc.add_paragraph(str_j)
                        option += str_j + "\n"
                answer = tab.s_ele('@class=right-ans').text.replace('\u2003', ':')
            elif topic == '不定项选择题':
                if tab.s_ele('xpath://*[@id="body"]/div/div[1]/div[2]/div[2]/div[3]/p[2]/span[2]/div').attr(
                        'aria-checked') is None:
                    try:
                        tab.ele('@@class^before-icon@@text()=A').click()
                        tab.wait(0.1)
                        tab.ele('xpath://*[@id="body"]/div/div[1]/div[2]/div[1]/div[1]/div/div[3]/button').click()
                    except Exception as e:
                        print(e)
                else:
                    pass
                options = tab.eles('@class^option')
                for j in options:
                    try:
                        x = j.s_ele('@class^before-icon')
                        option_img = j.ele('tag:img',timeout=0.1).get_screenshot(path=rf'.\{ID}\imgs\option',
                                                                                  name=f'ques{i + 1}-option-{x.text}.png')
                        para = doc.add_paragraph()
                        list_j = list(j.text)
                        list_j.insert(1, '.')
                        str_j = ''.join(list_j)
                        run = para.add_run(str_j)  # 添加选项文本
                        run.add_picture(option_img)
                    except Exception as e:
                        list_j = list(j.text)
                        list_j.insert(1, '.')
                        str_j = ''.join(list_j)
                        doc.add_paragraph(str_j)
                        option += str_j + "\n"
                answer = tab.s_ele('@class=right-ans').text.replace('\u2003', ':')
            elif topic == '排序题':
                options = tab.eles('@class^option')
                for j in options:
                    try:
                        x = j.s_ele('@class^before-icon')
                        option_img = j.ele('tag:img',timeout=0.1).get_screenshot(path=rf'.\{ID}\imgs\option',
                                                                                  name=f'ques{i + 1}-option-{x.text}.png')
                        para = doc.add_paragraph()
                        list_j = list(j.text)
                        list_j.insert(1, '.')
                        str_j = ''.join(list_j)
                        run = para.add_run(str_j)  # 添加选项文本
                        run.add_picture(option_img)
                    except Exception as e:
                        list_j = list(j.text)
                        list_j.insert(1, '.')
                        str_j = ''.join(list_j)
                        doc.add_paragraph(str_j)
                        option += str_j + "\n"
                answer = tab.s_ele('@class=right-ans').text.replace('\u2003', ':')
            elif topic == '填空题':
                try:
                    s = tab.s_ele('xpath://*[@id="body"]/div/div[1]/div[2]/div[2]/div[3]/p[2]/span[2]/div').attr(
                        'aria-checked')
                    if s == 'true':
                        pass
                    else:
                        try:
                            tab.ele('@class=el-input__inner').input('1')
                            tab.wait(0.1)
                            tab.ele(
                                'xpath://*[@id="body"]/div/div[1]/div[2]/div[1]/div[1]/div/div[2]/button').click()
                        except Exception as e:
                            print(e)
                except Exception as e:
                    print(e)
                answer = tab.s_ele('@class=right-ans').text.replace('\u2003', ':')
            elif topic == '简答题':
                try:
                    s = tab.s_ele('xpath://*[@id="body"]/div/div[1]/div[2]/div[2]/div[3]/p[2]/span[2]/div').attr(
                        'aria-checked')
                    if s == 'true':
                        pass
                    else:
                        try:
                            tab.ele('@class=el-textarea__inner').input('1')
                            tab.wait(0.1)
                            tab.ele(
                                'xpath://*[@id="body"]/div/div[1]/div[2]/div[1]/div[1]/div/div[2]/button[2]').click()
                        except Exception as e:
                            print(e)
                except Exception as e:
                    print(e)
                answer = tab.s_ele('@class=right-ans').text.replace('\u2003', ':')
            elif topic == '论述题':
                answer = tab.s_ele('@class=right-ans').text.replace('\u2003', ':')
            else:
                raise Exception(f'暂不支持{topic}')

            analysis = ''
            if anl_switch == '是':
                try:
                    if topic in ['论述题', '不定项选择题']:
                        analysis = '答案解析：' + tab.ele(
                            'xpath://*[@id="body"]/div/div[1]/div[2]/div[1]/div[2]/div[2]/p/p').text.replace('\n',
                                                                                                             '')
                    else:
                        analysis = tab.s_ele('@class=answer-analysis').text.replace('\n', '')
                    try:
                        analysis_img_attr = tab.ele('@class^answer-analysis').ele('tag:img',timeout=0.1)
                        if analysis_img_attr.link:
                            analysis_img_url = analysis_img_attr.attr('src')
                            if 'ai_tag.png' in analysis_img_url:
                                pass
                            else:
                                analysis_img = analysis_img_attr.get_screenshot(path=rf'.\{ID}\imgs\analysis',
                                                                                name=f'ques{i + 1}-analysis.png')
                    except Exception as e:
                        pass
                except Exception as e:
                    print(e)
            if option != '':
                ques = f'{i + 1}.{title}\n{option}{answer}\n{analysis}\n\n'
                option1 = option.replace('\n', '&@')
                ques1 = f'{i + 1}&@{title}&@{option1}&@{answer[5:]}&@{analysis}\n'
            else:
                ques = f'{i + 1}.{title}\n{option}{answer}\n{analysis}\n\n'
                ques1 = f'{i + 1}&@{title}&@{answer[5:]}&@{analysis}\n'
            # 添加答案段落
            doc.add_paragraph(answer)
            doc.add_paragraph(f'{analysis} \n')
            try:
                doc.add_picture(analysis_img)
            except Exception as e:
                pass
            list_a = ques1.split('&@')
            while len(list_a) <= 4:
                list_a.insert(2, '')
            while 4 < len(list_a) < 12:
                list_a.insert(-3, '')
            try:
                ws1.write(i + 1, 0, int(list_a[0]))
                ws1.write(i + 1, 1, list_a[1])
                ws1.write(i + 1, 2, list_a[2])
                ws1.write(i + 1, 3, list_a[3])
                ws1.write(i + 1, 4, list_a[4])
                ws1.write(i + 1, 5, list_a[5])
                ws1.write(i + 1, 6, list_a[6])
                ws1.write(i + 1, 7, list_a[7])
                ws1.write(i + 1, 8, list_a[8])
                ws1.write(i + 1, 9, list_a[9])
                ws1.write(i + 1, 10, list_a[-2])
                ws1.write(i + 1, 11, list_a[-1][5:])
            except IndexError as e:
                pass
            wb.save(rf'.\{ID}\{ID}-第{begin}题开始.xls')
            info = f'第{i + 1}题已完成'
            try:
                print(ques, flush=True)
            except Exception as e:
                print(e)
                print(info, flush=True)
            filepath = rf'.\{ID}\{ID}-第{begin}题开始.txt'
            with open(filepath, "a", encoding='utf8') as f:
                f.write(ques)  # 自带文件关闭功能，不需要再写f.close()
            doc.save(rf'.\{ID}\{ID}-第{begin}题开始.docx')
            try:
                tab.ele('@@class:el-button el-button--primary el-button--small@@text():下一题', timeout=5).click()
                tab.wait(float(delay))
            except Exception as e:
                print(e)
        except Exception as e:
            print(f'第{i + 1}题下载失败\n', flush=True)
            with open(f'{ID}_error_log.txt', "a", encoding='utf8') as f:
                f.write(f'第{i + 1}题下载失败，错误信息：{e}\n')  # 自带文件关闭功能，不需要再写f.close()

            try:
                tab.ele('@@class:el-button el-button--primary el-button--small@@text():下一题', timeout=5).click()
                tab.wait(float(delay))
            except Exception as e:
                print(e)
        continue
    if file_format == '.txt':
        os.startfile(rf'.\{ID}\{ID}-第{begin}题开始.txt')
    elif file_format == '.docx':
        os.startfile(rf'.\{ID}\{ID}-第{begin}题开始.docx')
    elif file_format == '.xls':
        os.startfile(rf'.\{ID}\{ID}-第{begin}题开始.xls')
    elif file_format == '不自动打开':
        print('不自动打开文件')
    try:
        os.startfile(f'{ID}_error_log.txt')
    except FileNotFoundError:
        print('全部完成,未生成错误日志')