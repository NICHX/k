from openpyxl import Workbook
import re
from DrissionPage import Chromium, ChromiumOptions
from DrissionPage.errors import ElementNotFoundError
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt
import os
import sys
import requests
import ddddocr
from PIL import Image, ImageDraw, ImageFont
from fontTools.ttLib import TTFont
from io import BytesIO
import logging
from DrissionPage.common import Settings


Settings.set_raise_when_ele_not_found(True)



def download_ques_advanced(question_id, name, delay, start_num, default_open, parse, timeout, chapter_id='', kid=''):
    try:
        os.makedirs(rf'.\{question_id}', exist_ok=True)
        print('目录创建成功', flush=True)
    except FileExistsError:
        print('目录已存在', flush=True)
        pass

    logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

    home_url = f'https://www.zaixiankaoshi.com/home/'
    login_url = f'https://www.zaixiankaoshi.com/login/'
    if chapter_id != '':
        url = f'https://www.zaixiankaoshi.com/online/?paperId={question_id}&chapter={chapter_id}&kid={kid}'
    else:
        url = f'https://www.zaixiankaoshi.com/online/?paperId={question_id}'

    co = ChromiumOptions()
    # co.incognito()
    # 连接浏览器
    browser = Chromium(co)
    # 获取标签页对象
    tab = browser.latest_tab
    if tab.url == home_url:
        tab = browser.new_tab()
    elif tab.url != home_url:
        tab.get(login_url)
        tab.wait.url_change('https://www.zaixiankaoshi.com/home/', timeout=120)
        tab = browser.new_tab()
    tab.get(url)
    tab.wait.eles_loaded('@class=topic-num', timeout=1)
    n = tab.ele('@class=topic-num', timeout=1).text
    number = re.findall(r"/(\d+)、", n)[0]

    # 打开背题模式
    try:
        button_on = tab.ele('@class=answer-box', timeout=float(timeout))
        if button_on:
            print('背题模式已打开')
    except ElementNotFoundError:
        print('尝试自动点击背题模式按钮')
        tab.eles('@class=el-switch__input',
                 timeout=1)[-1].click()
        tab.wait.eles_loaded('@class=answer-box', timeout=1)
        print('打开背题模式成功')
    except Exception as e:
        print('自动点击失败，请手动点击背题模式按钮后重试')
        tab.wait(10)

    if int(start_num) != 1:
        try:
            tab.ele(f'tag:span@text():{start_num}').click()
        except Exception as e:
            print(e)

    elif int(start_num) > int(number):
        print('起始题号超出题库范围！')
        sys.exit(1)

    def get_encrypt_ttf():
        ttf_attr = tab.ele('@class=qusetion-title', timeout=float(timeout)).attr('style')
        pattern = r':\s*(.*?);'
        ttf = re.search(pattern, ttf_attr).group(1)
        ttf_link = f'https://resource.zaixiankaoshi.com/fonts/{ttf}.ttf'
        ttf_file = tab.download(ttf_link, save_path=rf'.\{question_id}', file_exists='overwrite', show_msg=False)[1]
        print('ttf文件下载成功', flush=True)
        return ttf, ttf_file

    def convert_to_image(cmap_code, ttf):
        img_size = 1024
        img = Image.new(mode='1', size=(1024, 1024), color=255)
        draw = ImageDraw.Draw(img)
        font = ImageFont.truetype(ttf, int(1024 * 0.6))
        character = chr(cmap_code)
        bbox = draw.textbbox((0, 0), character, font=font)
        width = bbox[2] - bbox[0]
        height = bbox[3] - bbox[1]
        draw.text(((img_size - width) / 2, (img_size - height) / 5), text=character, font=font)
        return img

    def extract_text_from_ttf(ttf):
        font = TTFont(ttf)
        ocr = ddddocr.DdddOcr(show_ad=False)
        font_map = {}
        for cmap_code, glyph_name in font.getBestCmap().items():
            bytes_io = BytesIO()
            image = convert_to_image(cmap_code, ttf)
            image.save(bytes_io, format='PNG')
            ocr.set_ranges('7')
            text = ocr.classification(bytes_io.getvalue())
            # image.save(rf'./ocr/{text}.png', format='PNG')
            # print('text:',text)
            font_map[hex(cmap_code).replace('0x', '')] = text
        print('解密成功')
        return font_map

    def get_question_info(i, ttf, font_map, parse_s=parse):
        error_info = ''
        try:
            topic = tab.eles('@class=topic-type', timeout=float(timeout))
            case_dict = {}
            if len(topic) == 1:
                case_dict = {}
            elif len(topic) == 2:
                case_img_paths = []
                case = tab.ele('@class=parent_question_inner', timeout=float(timeout))
                case_text = f'{case.text}'.replace("\n", "")
                try:
                    case_img = case.eles('tag:img', timeout=float(timeout))
                    if case_img:
                        for j in case_img:
                            case_img_path = j.save(path=rf'.\{question_id}\case\{i + 1}', rename=False)
                            case_img_paths.append(case_img_path)
                    else:
                        pass
                except ElementNotFoundError:
                    pass

                def remove_case_parse(s):
                    if s.startswith("案例分析"):
                        return s[4:]  # 去掉前4个字符
                    return s
                case_text = remove_case_parse(case_text)
                case_dict = {case_text: case_img_paths}

            # 获取题干
            title = tab.ele('@class=qusetion-title', timeout=float(timeout))
            original_title = title.text.replace('\n', '')
            title_text = ''
            title_img_paths = []
            for char in original_title:
                char_code = hex(ord(char)).replace('0x', '')
                if char_code in font_map:
                    title_text += font_map[char_code]
                else:
                    title_text += char
            try:
                title_img = title.eles('tag:img', timeout=float(timeout))
                if title_img:
                    for j in title_img:
                        title_img_path = None
                        try:
                            title_img_path = j.save(path=rf'.\{question_id}\title\{i + 1}')
                            title_img_paths.append(title_img_path)
                        except ElementNotFoundError:
                            pass
                        except Exception as e:
                            print(f"保存图片时出错: {e}")
            except Exception as e:
                error_info = e
                pass

            title_dict = {f'{i + 1}.{title_text}': title_img_paths}

            # 获取选项
            def format_option_text(text):
                """
                格式化选项文字，在第二个字符后插入 '.'
                """
                option_list = list(text)
                option_list.insert(1, '.')
                return ''.join(option_list)

            option_dict = {}
            try:
                options = tab.eles('@class^option', timeout=float(timeout))
                for j in options:
                    option_img_path = None
                    try:
                        option_img = j.ele('tag:img', timeout=float(timeout))
                        if option_img:
                            option_img_path = option_img.save(path=rf'.\{question_id}\option\{i + 1}')
                            tab.wait(0.1)
                    except ElementNotFoundError:
                        pass
                    except Exception as e:
                        error_info = e
                        pass
                    option_text = format_option_text(j.text)
                    option_dict[option_text] = option_img_path
            except Exception as e:
                error_info = e
                print(f"获取选项元素时出现错误: {e}")

            # 获取答案
            if topic[-1].text in ['单选题', '多选题', '判断题', '不定项选择题', '排序题']:
                answer = tab.ele('@class=right-ans', timeout=float(timeout))
            else:
                answer = tab.ele('@class=mt20', timeout=float(timeout))
            answer_text = answer.text.replace('\u2003', ':')
            answer_img_paths = []
            try:
                answer_img = answer.eles('tag:img', timeout=float(timeout))
                if answer_img:
                    for j in answer_img:
                        answer_img_path = None
                        try:
                            answer_img_path = j.save(path=rf'.\{question_id}\answer\{i + 1}')
                            answer_img_paths.append(answer_img_path)
                        except Exception as e:
                            error_info = e
                            print(f"保存图片时出错: {e}")
            except ElementNotFoundError:
                pass
            except Exception as e:
                error_info = e
                pass
            answer_dict = {f'{answer_text.replace("正确答案:","")}': answer_img_paths}

            # 获取解析
            if parse_s == '是':
                parse_img_paths = []
                parse = tab.ele('@class^answer-analysis', timeout=float(timeout))
                parse_text = f'解析：{parse.text}'.replace("\n", "").rstrip("查看全部")
                try:
                    parse_img = parse.eles('tag:img', timeout=float(timeout))
                    if parse_img:
                        for j in parse_img:
                            parse_img_url = j.link
                            if 'ai_tag.png' in parse_img_url:
                                parse_text += '（ai解析）'
                            else:
                                for j in parse_img:
                                    parse_img_path = j.save(path=rf'.\{question_id}\parse\{i + 1}', rename=False)
                                    parse_img_paths.append(parse_img_path)
                    else:
                        pass
                except ElementNotFoundError:
                    pass
                parse_dict = {f'{parse_text}': parse_img_paths}
            else:
                parse_dict = {'':[]}
            return case_dict, title_dict, option_dict, answer_dict, parse_dict, True, error_info
        except Exception as e:
            error_info = e
            print(f"获取题目信息时出错: {e}")
            return None, None, None, None, None, None, None, False, error_info

    doc = Document()
    wb = Workbook()
    ws = wb.active
    headers = ['序号', '案例', '题目', 'A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', '正确答案', '解析']
    ws.append(headers)

    # 设置 DOCX 字体为宋体
    doc.styles['Normal'].font.name = '宋体'
    r = doc.styles['Normal']._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    ttf = get_encrypt_ttf()
    font_map = extract_text_from_ttf(ttf[1])
    for i in range(int(start_num) - 1, int(number)):
        max_retries = 3
        success = False
        ques = None
        for retry in range(max_retries):
            try:
                # 解密
                ttf_attr = tab.ele('@class=qusetion-title', timeout=float(timeout)).attr('style')
                pattern = r':\s*(.*?);'
                ttf_now = re.search(pattern, ttf_attr).group(1)
                if ttf_now == ttf[0]:
                    pass
                else:
                    ttf = get_encrypt_ttf()
                    font_map = extract_text_from_ttf(ttf[1])
                    print('解密更新成功')
            except Exception as e:
                print(e)
            ques = get_question_info(i, ttf, font_map)
            if ques[-2]:
                case_dict, title_dict, option_dict, answer_dict, parse_dict, _, _ = ques
                title_text = list(title_dict.keys())[0].split('.', 1)[-1]
                answer_text = list(answer_dict.keys())[0].replace('正确答案:', '')
                parse_text = list(parse_dict.keys())[0]

                option_values = [''] * 8
                option_letters = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']
                for option, _ in option_dict.items():
                    option_text = option.split('.', 1)[-1].strip()
                    letter = option[0]
                    if letter in option_letters:
                        index = option_letters.index(letter)
                        option_values[index] = option_text

                output = ''
                case_text = list(case_dict.keys())[0] if case_dict else ''
                if case_text:
                    output += f'案例: {case_text}\n'
                output += f'{i + 1}.{title_text}\n'
                for letter, option in zip(option_letters, option_values):
                    if option:
                        output += f'{letter}: {option}\n'
                output += f'正确答案:{answer_text}\n'
                output += f'{parse_text}\n'
                #output += '-' * 50 + '\n'
                print(output,flush=True)
                success = True
                break
            else:
                if retry < max_retries - 1:
                    print(f'第{i + 1}题下载失败，正在重试第{retry + 1}次...', flush=True)
                    tab.ele(f'tag:span@text():{i+1}').click()
                    tab.wait(float(delay), float(delay) + 1)
                else:
                    print(f'第{i + 1}题下载失败，已重试{max_retries}次', flush=True)
                    with open(rf'.\{question_id}\{question_id}_error_log.txt', "a", encoding='utf8') as f:
                        f.write(f'第{i + 1}题下载失败，报错信息：{ques[-1]}\n')  # 自带文件关闭功能，不需要再写f.close()
                    try:
                        tab.ele('@@class=tit@@text():下一题',
                                timeout=5).click()
                        tab.wait(float(delay) ,float(delay)+1)
                    except Exception as e:
                        pass
        if success:
            case_dict, title_dict, option_dict, answer_dict, parse_dict, _, _ = ques

            # 写入题目信息到 DOCX 文件
            if case_dict:
                case_text = list(case_dict.keys())[0]
                case_para = doc.add_paragraph()
                case_run = case_para.add_run('\n' + f'案例: {case_text}')
                case_run.font.name = '宋体'
                r = case_run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                case_run.bold = True  # 设置字体加粗
                case_run.font.size = Pt(14)  # 设置字号为 14pt
                case_img_paths = case_dict[case_text]
                for img_path in case_img_paths:
                    doc.add_picture(img_path, width=Inches(4))

            for title, title_img_paths in title_dict.items():
                heading = doc.add_heading(title, level=1)
                # 直接对 heading 这个 Paragraph 对象设置字体
                for run in heading.runs:
                    run.font.name = '宋体'
                    r = run._element
                    r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                for img_path in title_img_paths:
                    doc.add_picture(img_path, width=Inches(4))

            for option, option_img_path in option_dict.items():
                para = doc.add_paragraph(option)
                for run in para.runs:
                    run.font.name = '宋体'
                    r = run._element
                    r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                if option_img_path:
                    doc.add_picture(option_img_path, width=Inches(2))

            for answer, answer_img_paths in answer_dict.items():
                para = doc.add_paragraph(f'答案: {answer}')
                for run in para.runs:
                    run.font.name = '宋体'
                    r = run._element
                    r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                for img_path in answer_img_paths:
                    doc.add_picture(img_path, width=Inches(2))

            for parse, parse_img_paths in parse_dict.items():
                para = doc.add_paragraph(parse)
                for run in para.runs:
                    run.font.name = '宋体'
                    r = run._element
                    r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                for img_path in parse_img_paths:
                    doc.add_picture(img_path, width=Inches(2))

            # 写入题目信息到 XLSX 文件
            case_text = list(case_dict.keys())[0] if case_dict else ''
            # 去除题干中的题号
            title_text = list(title_dict.keys())[0].split('.', 1)[-1]
            answer_text = list(answer_dict.keys())[0].replace('正确答案:', '')
            parse_text = list(parse_dict.keys())[0].replace('解析：', '')

            option_values = [''] * 8
            option_letters = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']
            for option, _ in option_dict.items():
                # 去除选项中的字母和点号
                option_text = option.split('.', 1)[-1].strip()
                letter = option[0]
                if letter in option_letters:
                    index = option_letters.index(letter)
                    option_values[index] = option_text

            row = [i + 1, case_text, title_text] + option_values + [answer_text, parse_text]
            ws.append(row)

            # 写入题目信息到 TXT 文件
            with open(rf'.\{question_id}\{question_id}-{name}-第{start_num}题开始.txt', 'a', encoding='utf-8') as txt_file:
                if case_dict:
                    txt_file.write(f'案例: {case_text}\n')
                txt_file.write(f'{i + 1}.{title_text}\n')
                for letter, option in zip(option_letters, option_values):
                    if option:
                        txt_file.write(f'{letter}: {option}\n')
                txt_file.write(f'正确答案: {answer_text}\n')
                txt_file.write(f'解析:{parse_text}\n\n')
                #txt_file.write('-' * 50 + '\n')
            # 每下载一题保存一次文件
            doc.save(rf'.\{question_id}\{question_id}-{name}-第{start_num}题开始.docx')
            wb.save(rf'.\{question_id}\{question_id}-{name}-第{start_num}题开始.xlsx')

            try:
                tab.ele('@@class=tit@@text():下一题',
                        timeout=5).click()
                tab.wait(float(delay) ,float(delay)+1)
            except Exception as e:
                pass

    if default_open == '.txt':
        os.startfile(rf'.\{question_id}\{question_id}-{name}-第{start_num}题开始.txt')
    elif default_open == '.docx':
        os.startfile(rf'.\{question_id}\{question_id}-{name}-第{start_num}题开始.docx')
    elif default_open == '.xlsx':
        os.startfile(rf'.\{question_id}\{question_id}-{name}-第{start_num}题开始.xlsx')
    elif default_open == '不自动打开':
        print('不自动打开文件')


    def delete_ttf_files(directory):
        try:
            for root, _, files in os.walk(directory):
                for file in files:
                    if file.endswith('.ttf'):
                        file_path = os.path.join(root, file)
                        os.remove(file_path)
                        print(f"已删除: {file_path}")
        except Exception as e:
            print(f"发生错误: {e}")

    target_directory = '.'  # 当前目录
    delete_ttf_files(target_directory)

    try:
        os.startfile(rf'.\{question_id}\{question_id}_error_log.txt')
    except FileNotFoundError:
        print('全部完成,未生成错误日志')





