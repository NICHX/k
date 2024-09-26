# -*- coding: utf-8 -*-
import codecs
import os
import sys

from DrissionPage import ChromiumPage, SessionPage
from DrissionPage.common import Settings
from DrissionPage.errors import ElementNotFoundError
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt
import xlwt
from gooey import Gooey, GooeyParser

Settings.raise_when_ele_not_found = True

if sys.stdout.encoding != 'UTF-8':
    sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')
if sys.stderr.encoding != 'UTF-8':
    sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'strict')

version = '1.0.0'


@Gooey(language='chinese', program_name=u'kaoshibao工具(advanced_version)', required_cols=2, optional_cols=2,
       advanced=True, clear_before_run=True, sidebar_title='工具列表', terminal_font_family='Courier New',
       menu=[{
           'name': '关于',
           'items': [{
               'type': 'AboutDialog',
               'menuTitle': '关于',
               'name': 'kaoshibao工具(advanced_version)\n',
               'description': 'Created by NICHX !\n 1、可导出题库为TXT、Word、excel格式',

               'version': version,
           }]
       }])
def main_window():
    parser = GooeyParser(
        description="安装谷歌Chrome浏览器！加QQ群980236861获取口令")
    subs = parser.add_subparsers(help='考试宝', dest='command')
    normal_parser = subs.add_parser('考试宝', help='kaoshibao工具')
    subgroup = normal_parser.add_argument_group('配置信息')
    '''subgroup.add_argument('考试宝帐号', help="必填")
    subgroup.add_argument('考试宝密码', widget='PasswordField', help="必填")'''
    subgroup.add_argument('口令', help="请输入QQ群中获取的口令", widget='TextField')
    subgroup.add_argument('题库ID', help="请输入题库ID", widget='TextField')
    subgroup.add_argument('保存目录', help="请选择想要保存到的目录", widget='DirChooser')
    subgroup.add_argument('默认打开文件', help="请选择完成后想要打开的文件类型", widget='Dropdown'
                          , choices=['.txt', '.docx', '.xls', '不自动打开'], default='.docx')
    subgroup.add_argument('延迟时间', help="爬取延迟时间，默认0.4，若有题目重复手动调高", widget='TextField',
                          default='0.4')

    args = parser.parse_args()

    if args.command == '考试宝':
        download_ques(args.题库ID, args.保存目录, args.延迟时间, args.口令, args.默认打开文件)


def main():
    page = SessionPage()
    # 访问网页
    page.get('https://space.nichx.cn/advanced_version.txt')
    remote_version = page.ele('text:version').text[10:]
    if remote_version == version:
        print(f'当前版本为{version} , 是最新版本', flush=True)
        main_window()
    else:
        print(f'当前版本为{version} , 最新版本为{remote_version} , 请加入QQ群：980236861 获取最新版本', flush=True)
        input('Press Enter to exit...')


def download_ques(ID, path, time, code, format):
    page1 = SessionPage()
    page1.get('https://space.nichx.cn/code/code_advanced.txt')
    remote_code = page1.html
    if remote_code == code:
        print(f'正确口令为{remote_code} , 校验通过', flush=True)
    else:
        print(f'口令错误或已更新,请到QQ群获取最新正确口令或联系管理员', flush=True)
        sys.exit(1)

    doc = Document()
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = Pt(11)

    wb = xlwt.Workbook(encoding='utf-8')  # 新建一个excel文件
    ws1 = wb.add_sheet('sheet1')  # 添加一个新表，名字为first
    ws1.write(0, 0, '序号')
    ws1.write(0, 1, '题目')
    ws1.write(0, 2, 'A')
    ws1.write(0, 3, 'B')
    ws1.write(0, 4, 'C')
    ws1.write(0, 5, 'D')
    ws1.write(0, 6, 'E')
    ws1.write(0, 7, 'F')
    ws1.write(0, 8, 'G')
    ws1.write(0, 9, 'H')
    ws1.write(0, 10, '正确答案')
    ws1.write(0, 11, '解析')

    url = f'https://www.zaixiankaoshi.com/online/?paperId={ID}'
    page = ChromiumPage()
    page.get(url)
    page.wait.eles_loaded('xpath://*[@id="body"]/div[2]/div[1]/div[2]/div[1]/div/div[1]/div/div[1]/div/span[2]')
    number = page.ele('xpath://*[@id="body"]/div[2]/div[1]/div[2]/div[1]/div/div[1]/div/div[1]/div/span[2]').text[2:-1]
    # 打开背题模式
    try:
        button_off = page.s_ele('@@role=switch@@class=el-switch')
        if button_off:
            page.ele('xpath://*[@id="body"]/div[2]/div[1]/div[2]/div[2]/div[2]/div[1]/p[2]/span[2]/div/input').click()
            print('点击背题模式按钮')
            page.wait(0.3, 0.6)
    except ElementNotFoundError:
        print('背题模式已打开')
        page.wait(0.3)
    for i in range(int(number)):
        try:
            title = f"{page.ele('@class=qusetion-box').text}".replace('\n', '')
            doc.add_paragraph(f'{i + 1}.{title}')
            try:
                ques_img = page.s_ele('@class=qusetion-box').ele('tag:img')
                if ques_img.link:
                    ques_img_url = ques_img.attr('src')
                    # ques_img_url = f'{ques_img_url}'
                    ques_img = page.download(ques_img_url, rf'.\imgs\{ID}\ques', rename=f'ques{i + 1}-title',
                                             file_exists='skip')
                    page.wait(0.3)
                doc.add_picture(ques_img[1])
            except ElementNotFoundError:
                pass
            topic = page.ele('@class=topic-type').text
            option = ''
            if topic == '单选题':
                options = page.s_eles('@class^option')
                for j in options:
                    try:
                        # 定位选项内的类名以'fr-fic '开头的元素，假设它代表选项图片
                        option_img_url = j.s_ele('tag:img').link
                        # 定位当前选项内的类名以'before-icon'开头的元素
                        x = j.s_ele('@class^before-icon')  # 更改此处，确保使用当前选项的'before-icon'元素
                        # 下载选项图片到指定目录，并重命名
                        a = page.download(option_img_url, rf'.\imgs\{ID}\option', rename=f'ques{i + 1}-option-{x.text}',
                                          file_exists='skip')
                        page.wait(0.3)
                        para = doc.add_paragraph()
                        list_j = list(j.text)
                        list_j.insert(1, '.')
                        str_j = ''.join(list_j)
                        run = para.add_run(str_j)  # 添加选项文本
                        img_path = a[1]
                        run.add_picture(img_path, width=Inches(2.5))
                    except Exception as e:
                        list_j = list(j.text)
                        list_j.insert(1, '.')
                        str_j = ''.join(list_j)
                        doc.add_paragraph(str_j)
                        option += str_j + "\n"
                answer = page.s_ele('@class=right-ans').text.replace('\u2003', ':')[:-1]
            elif topic == '判断题':
                options = page.ele('@class^select-left').children('@class^option')
                for j in options:
                    list_j = list(j.text)
                    list_j.insert(1, '.')
                    str_j = ''.join(list_j)
                    doc.add_paragraph(str_j)
                    option += str_j + "\n"
                answer = page.s_ele('@class=right-ans').text.replace('\u2003', ':')[:-1]
            elif topic == '多选题':
                options = page.s_eles('@class^option')
                for j in options:
                    try:
                        # 定位选项内的类名以'fr-fic '开头的元素，假设它代表选项图片
                        option_img_url = j.s_ele('tag:img').link
                        # 定位当前选项内的类名以'before-icon'开头的元素
                        x = j.s_ele('@class^before-icon')  # 更改此处，确保使用当前选项的'before-icon'元素
                        # 下载选项图片到指定目录，并重命名
                        option_img = page.download(option_img_url, rf'.\imgs\{ID}\option',
                                                   rename=f'ques{i + 1}-option-{x.text}', file_exists='skip')
                        page.wait(0.3)
                        para = doc.add_paragraph()
                        list_j = list(j.text)
                        list_j.insert(1, '.')
                        str_j = ''.join(list_j)
                        run = para.add_run(str_j)  # 添加选项文本
                        img_path = option_img[1]
                        run.add_picture(img_path, width=Inches(2.5))
                    except Exception as e:
                        list_j = list(j.text)
                        list_j.insert(1, '.')
                        str_j = ''.join(list_j)
                        doc.add_paragraph(str_j)
                        option += str_j + "\n"
                answer = page.s_ele('@class=right-ans').text.replace('\u2003', ':')[:-1]
            elif topic == '填空题':
                answer = '正确答案:' + page.s_ele('@class=mt20').text.replace('\u2003', ':')
            elif topic == '简答题':
                answer = '正确答案:' + page.s_ele('@class=mt20').text.replace('\u2003', ':')

            '''formatted_option = "\n".join(
                f"{line[0]}. {line[1:]}" if line[0].isupper() else line for line in option.splitlines())'''

            try:
                analysis = page.s_ele('@class^answer-analysis').text.replace('\n', '')
                try:
                    analysis_img = page.s_ele('@class^answer-analysis').ele('tag:img')
                    if analysis_img.link:
                        analysis_img_url = analysis_img.attr('src')
                        if analysis_img_url == 'https://resource.zaixiankaoshi.com/mini/ai_tag.png':
                            pass
                        else:
                            analysis_img_url = f'{analysis_img_url}'
                            analysis_img = page.download(analysis_img_url, rf'.\imgs\{ID}\analysis',
                                                         rename=f'ques{i + 1}-analysis', file_exists='skip')
                            page.wait(0.3)
                except ElementNotFoundError:
                    pass
            except Exception as e:
                print(e)
            if option != '':
                ques = f'{i+1}.{title}\n{option}{answer}\n解析：{analysis}\n\n'
                option1 = option.replace('\n', '&@')
                ques1 = f'{i+1}&@{title}&@{option1}&@{answer[4:]}&@{analysis}\n'
            else:
                ques = f'{i+1}.{title}\n{option}{answer}\n解析：{analysis}\n\n'
                ques1 = f'{i+1}&@{title}&@{answer[4:]}&@{analysis}\n'
            # 添加答案段落
            doc.add_paragraph(answer)
            doc.add_paragraph(f'解析：{analysis} \n')
            try:
                doc.add_picture(analysis_img[1])
            except Exception as e:
                pass
            list_a = ques1.split('&@')
            while len(list_a) <= 4:
                list_a.insert(2, '')
            while 4 < len(list_a) < 12:
                list_a.insert(-3, '')
            try:
                ws1.write(i+1, 0, int(list_a[0]))
                ws1.write(i+1, 1, list_a[1])
                ws1.write(i+1, 2, list_a[2])
                ws1.write(i+1, 3, list_a[3])
                ws1.write(i+1, 4, list_a[4])
                ws1.write(i+1, 5, list_a[5])
                ws1.write(i+1, 6, list_a[6])
                ws1.write(i+1, 7, list_a[7])
                ws1.write(i+1, 8, list_a[8])
                ws1.write(i+1, 9, list_a[9])
                ws1.write(i+1, 10, list_a[-2])
                ws1.write(i+1, 11, list_a[-1])
            except IndexError as e:
                pass
            wb.save(f'{path}/{ID}.xls')
            info = f'第{i + 1}题已完成'
            print(info, flush=True)
            filepath = f'{path}/{ID}.txt'
            with open(filepath, "a", encoding='utf8') as f:
                f.write(ques)  # 自带文件关闭功能，不需要再写f.close()
            doc.save(f'{path}/{ID}.docx')
            try:
                page.ele('@@class:el-button el-button--primary el-button--small@@text():下一题', timeout=5).click()
                page.wait(float(time))
            except Exception as e:
                print(e)
        except ElementNotFoundError:
            print(f'第{i + 1}题下载失败\n', flush=True)
            with open('error_log.txt', "a", encoding='utf8') as f:
                f.write(f'第{i + 1}题下载失败\n')  # 自带文件关闭功能，不需要再写f.close()
            try:
                page.ele('@@class:el-button el-button--primary el-button--small@@text():下一题', timeout=5).click()
                page.wait(float(time))
            except Exception as e:
                print(e)
        continue
    if format == '.txt':
        os.startfile(f'{path}/{ID}.txt')
    elif format == '.docx':
        os.startfile(f'{path}/{ID}.docx')
    elif format == '.xls':
        os.startfile(f'{path}/{ID}.xls')
    elif format == '不自动打开':
        print('不自动打开文件')
    try:
        os.startfile('error_log.txt')
    except FileNotFoundError:
        print('全部完成,未生成错误日志')


if __name__ == '__main__':
    main()
